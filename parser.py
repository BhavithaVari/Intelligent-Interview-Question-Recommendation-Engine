"""
Extracts text from resumes (PDF/DOCX) and extracts skills
"""
import re
import pdfplumber
from docx import Document

# Skill vocabulary - you can add more skills here
SKILL_VOCAB = [
    "python", "django", "flask", "fastapi", "postgresql", "mysql", "mongodb",
    "redis", "docker", "kubernetes", "celery", "rest api", "graphql",
    "react", "node.js", "javascript", "typescript", "aws", "gcp", "azure",
    "git", "ci/cd", "linux", "sql", "pandas", "numpy", "machine learning",
    "tensorflow", "pytorch", "langchain", "java", "c++", "html", "css"
]

def extract_text_from_pdf(filepath: str) -> str:
    """Extract text from PDF file"""
    text_chunks = []
    with pdfplumber.open(filepath) as pdf:
        for page in pdf.pages:
            page_text = page.extract_text() or ""
            text_chunks.append(page_text)
    return "\n".join(text_chunks)

def extract_text_from_docx(filepath: str) -> str:
    """Extract text from DOCX file"""
    doc = Document(filepath)
    chunks = [p.text for p in doc.paragraphs]
    
    # Extract text from tables too
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                if cell.text.strip():
                    chunks.append(cell.text)
    
    return "\n".join(chunks)

def extract_text(filepath: str) -> str:
    """Extract text from either PDF or DOCX"""
    if filepath.lower().endswith(".pdf"):
        return extract_text_from_pdf(filepath)
    elif filepath.lower().endswith(".docx"):
        return extract_text_from_docx(filepath)
    else:
        raise ValueError(f"Unsupported file type: {filepath}")

def extract_skills(text: str) -> list[str]:
    """Find known skills in the text"""
    text_lower = text.lower()
    found = []
    for skill in sorted(SKILL_VOCAB, key=len, reverse=True):
        pattern = r"\b" + re.escape(skill) + r"\b"
        if re.search(pattern, text_lower):
            found.append(skill)
    return sorted(set(found))

def extract_years_experience(text: str) -> float | None:
    """Extract years of experience"""
    match = re.search(r"(\d+(?:\.\d+)?)\+?\s*years?", text, re.IGNORECASE)
    return float(match.group(1)) if match else None

def parse_resume(filepath: str) -> dict:
    """Parse resume and extract all information"""
    text = extract_text(filepath)
    return {
        "raw_text": text,
        "skills": extract_skills(text),
        "years_experience": extract_years_experience(text),
    }

def parse_job_description(jd_text: str) -> dict:
    """Parse job description and extract skills"""
    return {
        "raw_text": jd_text,
        "required_skills": extract_skills(jd_text),
    }